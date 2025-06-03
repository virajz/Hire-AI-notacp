from fastapi import APIRouter, File, UploadFile, HTTPException
from typing import Dict, Any
import io
import re

# PDF and DOCX libraries
import fitz  # PyMuPDF
import pdfplumber
import docx

router = APIRouter()

def extract_text_from_pdf(file_bytes: bytes) -> str:
    # Try PyMuPDF first
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        if text.strip():
            return text
    except Exception:
        pass
    # Fallback to pdfplumber
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            return text
    except Exception:
        pass
    return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = []
        # Extract paragraphs
        text.extend([p.text for p in doc.paragraphs if p.text.strip()])
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        return "\n".join(text)
    except Exception:
        return ""

def extract_fields(text: str) -> Dict[str, Any]:
    fields = {}
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    text_block = " ".join(lines)
    
    # Split into sections for better parsing
    sections = {}
    current_section = "HEADER"
    current_content = []
    
    for line in lines:
        if re.match(r'^[A-Z\s]{2,}$', line):  # Section headers are typically all caps
            if current_content:
                sections[current_section] = current_content
            current_section = line
            current_content = []
        else:
            current_content.append(line)
    if current_content:
        sections[current_section] = current_content

    # Name (look in first few non-empty lines, before email/phone)
    for line in lines[:5]:
        if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$', line) and not re.search(r'@|\d', line):
            fields["name"] = line
            break
    
    # Email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    fields["email"] = email_match.group(0) if email_match else ""
    
    # Phone
    phone_match = re.search(r'(\+?\d[\d\s\-]{7,}\d)', text)
    fields["phone"] = phone_match.group(0) if phone_match else ""
    
    # Location (typically found near the name at the top)
    location_patterns = [
        r'([A-Z][a-z]+(?:\s*,\s*(?:India|USA|UK|Canada)))\s*(?=\n|$)',  # City, Country
        r'([A-Z][a-z]+\s*,\s*[A-Z][a-z]+(?:\s*,\s*(?:India|USA|UK|Canada))?)',  # City, State, Country
        r'([A-Z][a-z]+\s*,\s*[A-Z]{2})',  # City, State code
    ]
    
    for pattern in location_patterns:
        for line in lines[:10]:  # Look in first 10 lines
            if re.search(r'@|\d|resume|cv|name|phone', line.lower()):  # Skip lines with contact info
                continue
            loc_match = re.search(pattern, line)
            if loc_match:
                potential_location = loc_match.group(1).strip()
                # Don't use the line if it contains the candidate's name
                if fields.get("name", "").split()[0] not in potential_location:
                    fields["location"] = potential_location
                    break
        if fields.get("location"):
            break
    
    # Current title from EXPERIENCE or SUMMARY section
    if "EXPERIENCE" in sections:
        exp_lines = sections["EXPERIENCE"]
        for line in exp_lines[:3]:  # Look at first few lines of experience
            if re.search(r'engineer|developer|architect', line.lower()):
                fields["current_title"] = line.split('-')[0].strip() if '-' in line else line.strip()
                break
    elif "SUMMARY" in sections:
        summary = " ".join(sections["SUMMARY"])
        title_match = re.search(r'(?:As\s+(?:an?|the)\s+)?([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Engineer|Developer|Architect))', summary)
        if title_match:
            fields["current_title"] = title_match.group(1)

    # Calculate years of experience from experience section dates
    total_months = 0
    if "EXPERIENCE" in sections:
        experience_lines = sections["EXPERIENCE"]
        date_pattern = re.compile(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})\s*(?:-|to)\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|Present)', re.IGNORECASE)
        
        date_sections = []
        current_company = None
        
        for line in experience_lines:
            # Skip bullet points and descriptions
            if line.strip().startswith('•') or line.strip().startswith('-'):
                continue
                
            date_match = date_pattern.search(line)
            if date_match:
                start_date = date_match.group(1).title()  # Normalize case
                end_date = date_match.group(2).title() if date_match.group(2).lower() != 'present' else 'May 2025'
                
                # Clean up any remaining full month names
                for month in ['January', 'February', 'March', 'April', 'June', 'July', 'August', 'September', 'October', 'November', 'December']:
                    start_date = start_date.replace(month, month[:3])
                    end_date = end_date.replace(month, month[:3])
                
                date_sections.append((start_date, end_date))
        
        # Calculate total experience avoiding overlaps
        if date_sections:
            # Sort by start date for proper calculation
            date_sections.sort(key=lambda x: convert_date_to_months(x[0]))
            
            # Calculate total duration for each position
            for start_date, end_date in date_sections:
                start_months = convert_date_to_months(start_date)
                end_months = convert_date_to_months(end_date)
                if start_months and end_months:
                    duration = end_months - start_months
                    if duration > 0:
                        total_months += duration
            
            for i in range(1, len(date_sections)):
                start = convert_date_to_months(date_sections[i][0])
                end = convert_date_to_months(date_sections[i][1])
                if start > current_end:  # No overlap
                    total_months += end - start
                elif end > current_end:  # Partial overlap
                    total_months += end - current_end
                current_end = max(current_end, end)
    
    fields["years_exp"] = str(total_months // 12) if total_months > 0 else ""
    
    # Hard skills with expanded categories and normalization mapping
    skill_patterns = {
        "languages": r'\b(Python|Go|C\+\+|Java|JavaScript|TypeScript|Ruby|Scala|Kotlin|Swift|Rust)\b',
        "ai_ml_libraries": r'\b(TensorFlow|PyTorch|Pandas|NumPy|Keras|Scikit-learn|SciPy)\b',
        "ai_ml_concepts": r'\b(Natural Language Processing|NLP|Computer Vision|Data Analysis|Machine Learning|Deep Learning|GenAI|Generative AI)\b',
        "web_frameworks": r'\b(Flask|Streamlit|FastAPI|Django|Spring Boot|Ruby on Rails|Node\.js|Express\.js)\b',
        "data_visualization": r'\b(Matplotlib|Seaborn|Plotly|Tableau|Power BI|Looker|D3\.js)\b',
        "cloud": r'\b(AWS|Amazon Web Services|GCP|Google Cloud Platform|Azure|Microsoft Azure|Lambda|EC2|S3|Cloud Functions|CloudFormation|ARM Templates|Azure DevOps)\b',
        "devops": r'\b(Docker|Kubernetes|Terraform|Jenkins|Git|CI/CD|GitLab|GitHub Actions|Ansible|Puppet|Chef)\b',
        "databases": r'\b(SQL|PostgreSQL|Postgres|MySQL|MongoDB|Redis|DynamoDB|Cassandra|SQLite|Oracle DB|Microsoft SQL Server)\b',
        "tools": r'\b(Linux|Nginx|Apache|Maven|Gradle|npm|yarn|Jupyter|ComfyUI|Airflow|Spark|Hadoop|Kafka)\b',
        "security": r'\b(OWASP|Burp Suite|ZAP|JMeter|IAM|Security|OAuth|Encryption|Firewalls|SIEM|SOAR)\b',
        "other_tech": r'\b(GraphQL|REST API|Microservices|Big Data|Data Engineering|ETL|Data Warehousing|Blockchain)\b' # Catch-all for other important tech
    }
    
    # Normalize skill names - ensure lowercase keys for matching
    skill_normalization = {
        'python': 'Python',
        'go': 'Go',
        'c++': 'C++',
        'java': 'Java',
        'javascript': 'JavaScript',
        'typescript': 'TypeScript',
        'tensorflow': 'TensorFlow',
        'pytorch': 'PyTorch',
        'pandas': 'Pandas',
        'numpy': 'NumPy',
        'keras': 'Keras',
        'scikit-learn': 'Scikit-learn',
        'scipy': 'SciPy',
        'natural language processing': 'NLP',
        'nlp': 'NLP',
        'computer vision': 'Computer Vision',
        'data analysis': 'Data Analysis',
        'machine learning': 'Machine Learning',
        'deep learning': 'Deep Learning',
        'genai': 'GenAI',
        'generative ai': 'GenAI',
        'flask': 'Flask',
        'streamlit': 'Streamlit',
        'fastapi': 'FastAPI',
        'django': 'Django',
        'spring boot': 'Spring Boot',
        'node.js': 'Node.js',
        'express.js': 'Express.js',
        'matplotlib': 'Matplotlib',
        'seaborn': 'Seaborn',
        'plotly': 'Plotly',
        'tableau': 'Tableau',
        'power bi': 'Power BI',
        'aws': 'AWS',
        'amazon web services': 'AWS',
        'gcp': 'GCP',
        'google cloud platform': 'GCP',
        'azure': 'Azure',
        'microsoft azure': 'Azure',
        'lambda': 'AWS Lambda', # Be more specific if possible or keep generic
        'ec2': 'EC2',
        's3': 'S3',
        'docker': 'Docker',
        'kubernetes': 'Kubernetes',
        'terraform': 'Terraform',
        'jenkins': 'Jenkins',
        'git': 'Git',
        'ci/cd': 'CI/CD',
        'gitlab': 'GitLab',
        'github actions': 'GitHub Actions',
        'sql': 'SQL',
        'postgresql': 'PostgreSQL',
        'postgres': 'PostgreSQL',
        'mysql': 'MySQL',
        'mongodb': 'MongoDB',
        'redis': 'Redis',
        'dynamodb': 'DynamoDB',
        'jupyter': 'Jupyter',
        'comfyui': 'ComfyUI',
        'owasp': 'OWASP',
        'security': 'Security Engineering', # Keep existing or make more specific
        'iam': 'IAM',
        'graphql': 'GraphQL',
        'rest api': 'REST API',
        'microservices': 'Microservices',
        # Add more normalizations as needed
    }
    
    all_skills = []
    seen_skills = set()  # Track normalized skills to avoid duplicates
    
    for category, pattern in skill_patterns.items():
        matches = re.finditer(pattern, text, re.I)
        for match in matches:
            skill = match.group(1)
            normalized_skill = skill_normalization.get(skill.lower(), skill.title())
            
            if normalized_skill.lower() not in seen_skills:
                all_skills.append(normalized_skill)
                seen_skills.add(normalized_skill.lower())
    
    # Sort skills alphabetically for consistency
    fields["hard_skills"] = sorted(all_skills)
    return fields

def convert_date_to_months(date_str: str) -> int:
    """Convert a date string like 'Jan 2020' to total months since year 0"""
    month_map = {
        'Jan': 1, 'Feb': 2, 'Mar': 3, 'Apr': 4, 'May': 5, 'Jun': 6,
        'Jul': 7, 'Aug': 8, 'Sep': 9, 'Oct': 10, 'Nov': 11, 'Dec': 12
    }
    try:
        parts = date_str.strip().split()
        if len(parts) != 2:
            return 0
        month, year = parts
        return int(year) * 12 + month_map.get(month[:3], 0)
    except (ValueError, KeyError, IndexError):
        return 0

@router.post("/parse_resume/")
async def parse_resume(file: UploadFile = File(...)):
    ext = file.filename.split(".")[-1].lower()
    file_bytes = await file.read()
    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
        if not text.strip():
            raise HTTPException(status_code=422, detail="No extractable text found in PDF. It may be scanned or image-based.")
    elif ext in ("doc", "docx"):
        text = extract_text_from_docx(file_bytes)
        if not text.strip():
            raise HTTPException(status_code=422, detail="No extractable text found in DOCX. The file may be corrupted or empty.")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF or DOCX.")
    fields = extract_fields(text)
    # Fallback to LLM if not enough fields (simulate for now)
    if not fields["name"] or not fields["current_title"]:
        fields["llm_fallback"] = True
        # Here you would call Groq API for LLM extraction
    else:
        fields["llm_fallback"] = False
    fields["raw_text"] = text[:1000]  # Truncate for demo
    return fields
